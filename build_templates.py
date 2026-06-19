"""Generate the two Word report templates for PIs to download from the dashboard.

Both templates include:
  - Detailed milestone-tracking tables (with planned vs actual dates, status, outputs)
  - Line-item budget tables (approved / spent / balance / % utilised)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0B, 0x25, 0x45)
PURPLE = RGBColor(0x5B, 0x21, 0xB6)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
RED = RGBColor(0xDC, 0x26, 0x26)


def make_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
    return doc


def add_header_block(doc, title_text, kind):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MANILA CENTRAL UNIVERSITY")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = PURPLE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Institutional Research Office")
    r.italic = True; r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("In-House Research Grant — " + title_text)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{kind} Report Template")
    r.italic = True; r.font.size = Pt(11)
    doc.add_paragraph()


def add_h(doc, text, color=NAVY, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color


def add_field(doc, label, hint=""):
    p = doc.add_paragraph()
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(f"[ {hint} ]" if hint else "[ ]").italic = True


def add_textbox(doc, label, lines=4):
    add_h(doc, label, size=11)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    for _ in range(lines):
        cell.add_paragraph()


def add_table(doc, headers, n_rows=4, header_color=None):
    table = doc.add_table(rows=n_rows + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10)
        if header_color:
            r.font.color.rgb = header_color
    return table


def add_legend(doc, items):
    """Small italic legend below a table."""
    p = doc.add_paragraph()
    r = p.add_run("Legend: " + " · ".join(items))
    r.italic = True; r.font.size = Pt(9)


# ============================================================
# STANDARD BUDGET-CATEGORY ROWS — used in both templates
# ============================================================
BUDGET_CATEGORIES = [
    "Personnel / Honoraria",
    "Supplies & Reagents",
    "Equipment",
    "Travel & Fieldwork",
    "Participant Incentives",
    "Publication / Page Charges",
    "Dissemination / Conference",
    "Overhead / Administrative",
    "Contingency",
]


def add_budget_table(doc, columns, n_extra_rows=2):
    """Build budget table with the 9 standard categories pre-populated + extra blank rows."""
    n_rows = len(BUDGET_CATEGORIES) + n_extra_rows + 1   # +1 for total
    table = doc.add_table(rows=n_rows + 1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10)
    # Pre-populate category column
    for i, cat in enumerate(BUDGET_CATEGORIES, start=1):
        table.rows[i].cells[0].text = cat
    # "Other" rows
    for i in range(len(BUDGET_CATEGORIES) + 1, n_rows):
        table.rows[i].cells[0].text = "Other (specify)"
    # Total row
    total_cell = table.rows[n_rows].cells[0]
    total_cell.text = ""
    r = total_cell.paragraphs[0].add_run("TOTAL")
    r.bold = True
    return table


# ============================================================
# MID-PROJECT REPORT TEMPLATE
# ============================================================
def build_mid_report():
    doc = make_doc()
    add_header_block(doc, "Progress Report", "Mid-Project")

    # ---- 1. Project Info ----
    add_h(doc, "1. Project Information")
    add_field(doc, "Grant / Project ID", "e.g., CAS01")
    add_field(doc, "Project Title")
    add_field(doc, "Principal Investigator")
    add_field(doc, "Co-Investigators / Team")
    add_field(doc, "College / Department")
    add_field(doc, "Awarded Budget (₱)")
    add_field(doc, "Project Duration (months)")
    add_field(doc, "Start Date")
    add_field(doc, "Mid-Point Date")
    add_field(doc, "Reporting Period (From – To)")

    # ---- 2. Executive Summary ----
    add_textbox(doc, "2. Executive Summary (200–300 words)", lines=6)

    # ---- 3. Status of Objectives ----
    add_textbox(doc, "3. Status of Specific Objectives — what has been accomplished against each objective",
                lines=8)

    # ---- 4. MILESTONES (detailed table) ----
    add_h(doc, "4. Milestones — Mid-Point Status", color=PURPLE)
    p = doc.add_paragraph()
    r = p.add_run(
        "List every project milestone with its planned target date and the current status. "
        "Add or remove rows as needed."
    )
    r.italic = True; r.font.size = Pt(10)

    add_table(doc,
              ["#", "Milestone Description", "Planned Date",
               "Status", "Actual / Expected Date", "Outputs / Evidence", "Comments"],
              n_rows=8)
    add_legend(doc, [
        "Status: ✅ Completed", "🟡 In Progress", "🔴 Delayed", "⚪ Not Started",
    ])
    doc.add_paragraph()

    # ---- 5. Methodology Updates ----
    add_textbox(doc, "5. Methodology Updates / Protocol Amendments", lines=5)

    # ---- 6. Preliminary Results ----
    add_textbox(doc, "6. Preliminary Results / Findings to Date", lines=8)

    # ---- 7. BUDGET TRACKING ----
    add_h(doc, "7. Budget Utilisation — Line-Item Tracking", color=PURPLE)
    p = doc.add_paragraph()
    r = p.add_run(
        "Report cumulative spending against each approved budget category. "
        "Include % utilisation and balance remaining. Categories with > 10% over-budget "
        "must be justified in Section 8."
    )
    r.italic = True; r.font.size = Pt(10)

    add_budget_table(doc,
                     ["Category", "Approved (₱)", "Spent to Date (₱)",
                      "Encumbered (₱)", "Balance (₱)", "% Used", "Notes"])
    doc.add_paragraph()

    # ---- 7b. Quarterly snapshot ----
    add_h(doc, "7a. Quarterly Spending Snapshot", color=PURPLE, size=11)
    add_table(doc,
              ["Quarter", "Planned Spend (₱)", "Actual Spend (₱)",
               "Variance (₱)", "Cumulative Spend (₱)", "Notes"],
              n_rows=4)
    doc.add_paragraph()

    # ---- 8. Variances & Mitigation ----
    add_textbox(doc, "8. Budget Variances & Justifications — explain any over- or under-spending > 10% per category",
                lines=5)

    # ---- 9. Challenges ----
    add_textbox(doc, "9. Challenges Encountered & Mitigation", lines=5)

    # ---- 10. Dissemination ----
    add_h(doc, "10. Dissemination Activities to Date", color=PURPLE)
    add_table(doc,
              ["Type (Journal / Conf / Poster / Brief / Workshop)",
               "Title", "Venue / Journal", "Date", "Status"],
              n_rows=4)
    doc.add_paragraph()

    # ---- 11. Next Steps ----
    add_h(doc, "11. Plan for the Remaining Project Period", color=PURPLE)
    add_table(doc,
              ["Period", "Planned Activity / Milestone", "Owner",
               "Target Date", "Expected Output"],
              n_rows=6)
    doc.add_paragraph()

    # ---- 12. Ethics ----
    add_h(doc, "12. Ethics & Compliance")
    add_field(doc, "ERB Protocol Number")
    add_field(doc, "ERB Approval Date")
    add_field(doc, "Any protocol amendments since approval?",
              "Yes / No — describe if Yes")
    add_field(doc, "Adverse events or deviations to report?",
              "Yes / No — describe if Yes")

    # ---- 13. Signatures ----
    add_h(doc, "13. Certification & Signature")
    add_field(doc, "Principal Investigator (Name & Signature)")
    add_field(doc, "Date")
    add_field(doc, "Endorsed by Dean / Head of College (Name & Signature)")
    add_field(doc, "Date")

    return doc


# ============================================================
# FINAL REPORT TEMPLATE
# ============================================================
def build_final_report():
    doc = make_doc()
    add_header_block(doc, "Final Report", "Final Project")

    # ---- 1. Project Info ----
    add_h(doc, "1. Project Information")
    add_field(doc, "Grant / Project ID")
    add_field(doc, "Project Title")
    add_field(doc, "Principal Investigator")
    add_field(doc, "Co-Investigators / Team")
    add_field(doc, "College / Department")
    add_field(doc, "Awarded Budget (₱)")
    add_field(doc, "Planned Duration (months)")
    add_field(doc, "Actual Duration (months)")
    add_field(doc, "Start Date")
    add_field(doc, "End Date")

    # ---- 2. Abstract ----
    add_textbox(doc, "2. Abstract (250–400 words) — what was done, key findings, significance",
                lines=10)

    # ---- 3. Objectives & Achievement ----
    add_textbox(doc, "3. Objectives & Achievement — for each specific objective, describe what was achieved vs planned",
                lines=10)

    # ---- 4. MILESTONES — FINAL STATUS ----
    add_h(doc, "4. Milestones — Final Status", color=PURPLE)
    p = doc.add_paragraph()
    r = p.add_run(
        "Report the final status of every planned milestone. Mark each as Completed, "
        "Partially Completed, or Not Met, with evidence and explanation."
    )
    r.italic = True; r.font.size = Pt(10)

    add_table(doc,
              ["#", "Milestone Description", "Planned Date", "Actual Date",
               "Final Status", "Output / Evidence", "Explanation if Not Met"],
              n_rows=8)
    add_legend(doc, [
        "Final Status: ✅ Completed", "🟡 Partially Completed", "🔴 Not Met",
    ])
    doc.add_paragraph()

    # ---- 5. Methodology ----
    add_textbox(doc, "5. Methodology Summary — materials, methods, sample/data sources, analytical approach",
                lines=10)

    # ---- 6. Results ----
    add_textbox(doc, "6. Key Results & Findings", lines=14)

    # ---- 7. Discussion ----
    add_textbox(doc, "7. Discussion & Implications", lines=10)

    # ---- 8. Conclusions ----
    add_textbox(doc, "8. Conclusions & Recommendations", lines=8)

    # ---- 9. BUDGET TRACKING — FINAL ----
    add_h(doc, "9. Budget Utilisation — Final Reconciliation", color=PURPLE)
    p = doc.add_paragraph()
    r = p.add_run(
        "Report final spending per category. Provide written justification for "
        "any category with > 10% variance from the approved budget. Attach receipts "
        "or vouchers as a separate Annex."
    )
    r.italic = True; r.font.size = Pt(10)

    add_budget_table(doc,
                     ["Category", "Approved (₱)", "Spent (₱)",
                      "Variance (₱)", "% Variance", "Justification"])
    doc.add_paragraph()

    # ---- 9a. Quarterly Spend Summary ----
    add_h(doc, "9a. Quarterly Spending Summary", color=PURPLE, size=11)
    add_table(doc,
              ["Quarter", "Planned Spend (₱)", "Actual Spend (₱)",
               "Cumulative Spend (₱)", "% of Total Budget"],
              n_rows=5)
    doc.add_paragraph()

    # ---- 9b. Returned/Unspent Funds ----
    add_h(doc, "9b. Unspent Funds", color=PURPLE, size=11)
    add_field(doc, "Total Unspent (₱)")
    add_field(doc, "Reason for Unspent Balance")
    add_field(doc, "Status of Return to IRO")

    # ---- 10. Deliverables ----
    add_h(doc, "10. Final Deliverables Status", color=PURPLE)
    add_table(doc,
              ["Deliverable", "Planned", "Delivered",
               "Status (Completed / Partial / Not Met)", "Evidence / File / Reference"],
              n_rows=6)
    doc.add_paragraph()

    # ---- 11. Publications ----
    add_h(doc, "11. Publications & Dissemination", color=PURPLE)
    add_table(doc,
              ["Type (Journal / Conf / Poster / Brief / Book Chapter)",
               "Title", "Venue / Journal", "Year", "DOI / Link", "MCU Acknowledged?"],
              n_rows=5)
    doc.add_paragraph()

    # ---- 12. Impact ----
    add_textbox(doc, "12. Impact on MCU Research Agenda & SDGs", lines=4)

    # ---- 13. Limitations ----
    add_textbox(doc, "13. Limitations of the Study", lines=4)

    # ---- 14. Future Work ----
    add_textbox(doc, "14. Future Work Arising from this Project", lines=4)

    # ---- 15. Ethics ----
    add_h(doc, "15. Ethics & Compliance")
    add_field(doc, "ERB Protocol Number")
    add_field(doc, "ERB Approval Date")
    add_field(doc, "Protocol amendments during project?",
              "Yes / No — describe if Yes")
    add_field(doc, "Adverse events / serious deviations?",
              "Yes / No — describe if Yes")
    add_field(doc, "Data Availability — where will data be deposited?")

    # ---- 16. Funding acknowledgment ----
    add_h(doc, "16. Funding Acknowledgment Statement")
    p = doc.add_paragraph()
    r = p.add_run("All resulting publications and presentations must include: ")
    r.font.size = Pt(10); r.italic = True
    p = doc.add_paragraph()
    r = p.add_run(
        "\"This research was funded by the Manila Central University "
        "In-House Research Grant (Project ID: ___).\""
    )
    r.font.size = Pt(10); r.bold = True

    # ---- 17. Signatures ----
    add_h(doc, "17. Certification & Signature")
    add_field(doc, "Principal Investigator (Name & Signature)")
    add_field(doc, "Date")
    add_field(doc, "Endorsed by Dean / Head of College (Name & Signature)")
    add_field(doc, "Date")
    add_field(doc, "Received & Reviewed by IRO (office use only)")
    add_field(doc, "Date Received")

    return doc


# ============================================================
# WRITE FILES
# ============================================================
mid_path = "templates/MCU_Mid_Project_Report_Template.docx"
final_path = "templates/MCU_Final_Project_Report_Template.docx"
build_mid_report().save(mid_path)
build_final_report().save(final_path)
print(f"✅ Saved {mid_path}")
print(f"✅ Saved {final_path}")
